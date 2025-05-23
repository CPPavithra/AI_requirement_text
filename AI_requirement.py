# -*- coding: utf-8 -*-
"""Untitled6.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1r4MEuMgKXAqvm9DYphx9YN_G2lb5XsK9
"""

# Install necessary libraries
!pip install pyngrok pdfplumber python-docx openpyxl gradio requests

# Download ngrok (for Google Colab)
!wget https://bin.equinox.io/c/4VmDzA7iaHb/ngrok-stable-linux-amd64.zip -O ngrok.zip
!unzip -o ngrok.zip
!chmod +x ngrok

# Set the ngrok executable path
import os
from pyngrok import ngrok, conf

ngrok_executable = "./ngrok"
conf.get_default().ngrok_path = ngrok_executable
print(f"ngrok executable path set to: {conf.get_default().ngrok_path}")

# Set the authtoken directly using the ngrok client (modern command)
authtoken = "YOUR_AUTHTOKEN"  # REPLACE WITH YOUR ACTUAL AUTHTOKEN
!./ngrok authtoken $authtoken
print(f"Attempted to set authtoken via ngrok client: {authtoken}")

# Try to disconnect any existing ngrok tunnels
try:
    ngrok.disconnect()
    print("Successfully disconnected any existing ngrok tunnels.")
except Exception as e:
    print(f"Error disconnecting ngrok tunnels: {e}")

# Now you can continue with the rest of the setup
import pdfplumber
from docx import Document
import openpyxl
import email
from email import policy
import pandas as pd
import gradio as gr
import requests

# External knowledge integration (e.g., public regulations, standards)
def integrate_external_knowledge(extracted_text):
    standards = {
        "security": "Ensure data encryption for all communications.",
        "usability": "System should have a user-friendly interface."
    }
    extracted_text += "\n\n" + "Standards applied:\n"
    extracted_text += "\n".join([f"{key}: {value}" for key, value in standards.items()])
    return extracted_text

# Extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        return f"Error extracting text from PDF: {e}"
    return text.strip()

# Extract text from DOCX
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"
    return text.strip()

# Extract text from XLSX
def extract_text_from_excel(excel_path):
    text = ""
    try:
        workbook = openpyxl.load_workbook(excel_path)
        for sheet in workbook.sheetnames:
            ws = workbook[sheet]
            for row in ws.iter_rows(values_only=True):
                text += " | ".join(map(str, row)) + "\n"
    except Exception as e:
        return f"Error extracting text from XLSX: {e}"
    return text.strip()

# Extract text from EML (emails)
def extract_text_from_eml(eml_path):
    try:
        with open(eml_path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
            try:
                return msg.get_payload(decode=True).decode()
            except AttributeError:
                return str(msg.get_payload()) # Handle cases where get_payload is not a list
    except Exception as e:
        return f"Error extracting text from EML: {e}"
    return text.strip()

# Function to prompt users for further clarification (simulated here)
def ask_for_clarification(extracted_text):
    return extracted_text + "\n\nWhat is the priority level of this requirement?"

def process_file(file_path): # Changed uploaded_file to file_path
    if file_path is None:
        return "No file uploaded.", None, None

    filename = os.path.basename(file_path) # Extract filename from the path
    ext = filename.split('.')[-1].lower()

    extracted_text = ""
    if ext == "pdf":
        extracted_text = extract_text_from_pdf(file_path) # Use the path directly
    elif ext == "docx":
        extracted_text = extract_text_from_docx(file_path)
    elif ext == "xlsx":
        extracted_text = extract_text_from_excel(file_path)
    elif ext == "eml":
        extracted_text = extract_text_from_eml(file_path)
    else:
        return "Unsupported file type.", None, None

    if extracted_text.startswith("Error extracting text"):
        return extracted_text, None, None

    # Integrate external knowledge (standards, regulations, etc.)
    extracted_text = integrate_external_knowledge(extracted_text)

    # Ask the user for clarifications (simulating real-time questions)
    clarified_text = ask_for_clarification(extracted_text)

    # Generate the requirement document (Word format)
    try:
        doc = Document()
        doc.add_paragraph(clarified_text)
        doc.save("requirements.docx")
        requirements_file = "requirements.docx"
    except Exception as e:
        requirements_file = f"Error generating requirements.docx: {e}"

    # Generate user stories and store them in Excel
    try:
        user_stories = [
            {"User Story": "As a user, I want a secure login process.", "Priority": "High"},
            {"User Story": "As a user, I want an intuitive dashboard.", "Priority": "Medium"}
        ]
        df = pd.DataFrame(user_stories)
        df.to_excel("user_stories.xlsx", index=False)
        user_stories_file = "user_stories.xlsx"
    except Exception as e:
        user_stories_file = f"Error generating user_stories.xlsx: {e}"

    return "Processing complete! Download results:", requirements_file, user_stories_file

# Gradio UI for file uploads and displaying results
def gradio_ui():
    with gr.Blocks(title="AI-powered Requirement Extraction System") as iface:
        gr.Markdown("# AI-powered Requirement Extraction System")
        gr.Markdown("Upload a document (PDF, Word, Excel, Email) and extract and process functional & non-functional requirements.")

        with gr.Row():
            input_file = gr.File(label="Upload Document")

        with gr.Row():
            output_text = gr.Textbox(label="Processing Status", placeholder="Processing will start after you upload and submit a file.")

        with gr.Row():
            output_requirements = gr.File(label="Download Requirements (docx)", interactive=False)
            output_user_stories = gr.File(label="Download User Stories (xlsx)", interactive=False)

        input_file.upload(
            fn=process_file,
            inputs=input_file,
            outputs=[output_text, output_requirements, output_user_stories]
        )
    return iface.launch(server_port=7860, share=True, debug=True) # Added debug=True

# Start the Gradio UI and set up ngrok tunnel for public access
if __name__ == "__main__":
    app = gradio_ui()
